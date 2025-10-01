from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string, get_template
from django.utils.text import slugify
from django.utils.html import strip_tags
from django.conf import settings
from weasyprint import HTML
from django.core.files.base import ContentFile
from io import BytesIO
from docx import Document
from htmldocx import HtmlToDocx
from admisiones2.utils import generar_texto_comidas
from django.db import transaction
import logging
import tempfile

logger = logging.getLogger("admisiones2.services.informes")

from admisiones2.models.admisiones import (
    InformeTecnico,
    CampoASubsanar,
    ObservacionGeneralInforme,
    InformeTecnicoPDF,
    Admision,
    InformeComplementario,
    InformeComplementarioCampos,
    Anexo,
)
from admisiones2.forms.admisiones_forms import (
    InformeTecnicoJuridicoForm,
    InformeTecnicoBaseForm,
    AnexoForm,
)


class InformeService:
    @staticmethod
    def get_form_class_por_tipo(tipo):
        try:
            return (
                InformeTecnicoJuridicoForm
                if tipo == "juridico"
                else InformeTecnicoBaseForm
            )
        except Exception:
            logger.exception(
                "Error en get_form_class_por_tipo",
                extra={"tipo": tipo},
            )
            return InformeTecnicoBaseForm

    @staticmethod
    def get_tipo_from_kwargs(kwargs):
        try:
            return kwargs.get("tipo", "base")
        except Exception:
            logger.exception(
                "Error en get_tipo_from_kwargs",
                extra={"kwargs": kwargs},
            )
            return "base"

    @staticmethod
    def get_queryset_informe_por_tipo(tipo):
        try:
            return InformeTecnico.objects.filter(tipo=tipo)
        except Exception:
            logger.exception(
                "Error en get_queryset_informe_por_tipo",
                extra={"tipo": tipo},
            )
            return InformeTecnico.objects.none()

    @staticmethod
    def get_admision_y_tipo_from_kwargs(kwargs):
        try:
            tipo = kwargs.get("tipo", "base")
            admision_id = kwargs.get("admision_id")
            admision = get_object_or_404(Admision, pk=admision_id)
            return admision, tipo
        except Exception:
            logger.exception(
                "Error en get_admision_y_tipo_from_kwargs",
                extra={"kwargs": kwargs},
            )
            return None, "base"

    @staticmethod
    def get_anexo_form(admision, data=None, files=None, require_full=False):
        """Retorna un formulario de Anexo vinculado a la admisión."""
        try:
            anexo_instance = Anexo.objects.filter(admision=admision).last()
            if data is not None or files is not None:
                return AnexoForm(
                    data,
                    files,
                    instance=anexo_instance,
                    admision=admision,
                    require_full=require_full,
                )
            return AnexoForm(
                instance=anexo_instance, admision=admision, require_full=require_full
            )
        except Exception:
            logger.exception(
                "Error en get_anexo_form",
                extra={"admision_pk": getattr(admision, "pk", admision)},
            )
            if data is not None or files is not None:
                return AnexoForm(
                    data, files, admision=admision, require_full=require_full
                )
            return AnexoForm(admision=admision, require_full=require_full)

    @staticmethod
    def verificar_estado_para_revision(informe, action=None):
        """Actualiza estado al modificar un informe existente según la acción."""
        try:
            if action == "draft":
                informe.estado_formulario = "borrador"
                informe.estado = "Iniciado"
            else:
                if informe.estado != "Validado":
                    CampoASubsanar.objects.filter(informe=informe).delete()
                    ObservacionGeneralInforme.objects.filter(informe=informe).delete()
                    informe.estado_formulario = "finalizado"
                    informe.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en verificar_estado_para_revision",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )

    @staticmethod
    def get_campos_visibles_informe(informe):
        try:
            campos_excluidos_comunes = ["id", "admision", "estado", "tipo"]

            if informe.tipo == "juridico":
                campos_excluidos_especificos = [
                    "declaracion_jurada_recepcion_subsidios",
                    "constancia_inexistencia_percepcion_otros_subsidios",
                    "organizacion_avalista_1",
                    "organizacion_avalista_2",
                    "material_difusion_vinculado",
                ]
            elif informe.tipo == "base":
                campos_excluidos_especificos = [
                    "validacion_registro_nacional",
                    "IF_relevamiento_territorial",
                ]
            else:
                campos_excluidos_especificos = []

            campos_excluidos = campos_excluidos_comunes + campos_excluidos_especificos

            return [
                (field.verbose_name, getattr(informe, field.name))
                for field in informe._meta.fields
                if field.name not in campos_excluidos
            ]
        except Exception:
            logger.exception(
                "Error en get_campos_visibles_informe",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return []

    @staticmethod
    def preparar_informe_para_creacion(instance, admision_id, action=None):
        """Inicializa un informe técnico nuevo según la acción (borrador/finalizado)."""
        try:
            instance.admision_id = admision_id
            if action == "draft":
                instance.estado_formulario = "borrador"
                instance.estado = "Iniciado"
            else:
                instance.estado_formulario = "finalizado"
                instance.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en preparar_informe_para_creacion",
                extra={"admision_pk": admision_id},
            )

    @staticmethod
    def get_informe_por_tipo_y_pk(tipo, pk):
        try:
            return get_object_or_404(InformeTecnico, tipo=tipo, pk=pk)
        except Exception:
            logger.exception(
                "Error en get_informe_por_tipo_y_pk",
                extra={"tipo": tipo, "informe_pk": pk},
            )
            return None

    @staticmethod
    def actualizar_estado_informe(informe, nuevo_estado, tipo=None):
        try:
            informe.estado = nuevo_estado
            informe.save()

            if nuevo_estado == "Validado":
                InformeService.generar_y_guardar_pdf(informe, tipo)
        except Exception:
            logger.exception(
                "Error en actualizar_estado_informe",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )

    @staticmethod
    def generar_y_guardar_pdf(informe, tipo):
        """
        Genera y guarda:
          - PDF desde 'admisiones2/pdf/pdf_informe_tecnico.html' (WeasyPrint)
          - DOCX desde 'admisiones2/pdf/docx_informe_tecnico.html' (htmldocx)
        Si el template DOCX no existe, usa el mismo HTML del PDF.
        """
        try:
            anexo = Anexo.objects.filter(admision=informe.admision).first()
            try:
                texto_comidas = generar_texto_comidas(anexo) if anexo else {}
            except Exception as e:
                logger.warning("generar_texto_comidas falló: %s", e, exc_info=True)
                texto_comidas = {}

            context = {
                "informe": informe,
                "anexo": anexo,
                "texto_comidas": texto_comidas,
            }

            pdf_template_name = "admisiones2/pdf_informe_tecnico.html"
            html_pdf = render_to_string(pdf_template_name, context)
            if not html_pdf.strip():
                raise ValueError(
                    f"El template {pdf_template_name} devolvió contenido vacío."
                )

            docx_template_name = "admisiones2/docx_informe_tecnico.html"
            try:
                get_template(docx_template_name)
                html_docx = render_to_string(docx_template_name, context)
                if not html_docx.strip():
                    html_docx = html_pdf
            except Exception:
                html_docx = html_pdf

            base_url = str(
                getattr(settings, "STATIC_ROOT", "")
                or getattr(settings, "BASE_DIR", "")
                or "."
            )
            pdf_bytes = HTML(string=html_pdf, base_url=base_url).write_pdf()
            if not pdf_bytes:
                raise ValueError(
                    "WeasyPrint no devolvió contenido para el PDF generado."
                )

            docx_content = None
            try:
                doc = Document()
                HtmlToDocx().add_html_to_document(html_docx, doc)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                docx_bytes = buffer.getvalue()
                buffer.close()

                if not docx_bytes:
                    raise ValueError("htmldocx/python-docx devolvió contenido vacío.")
                docx_content = ContentFile(docx_bytes, name="tmp.docx")

            except Exception as docx_exc:
                logger.warning(
                    "Falla al convertir HTML a DOCX con htmldocx; se intentará un fallback simplificado.",
                    exc_info=docx_exc,
                    extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
                )
                try:
                    fallback_doc = Document()
                    fallback_text = strip_tags(html_docx)
                    for line in filter(
                        None,
                        (segment.strip() for segment in fallback_text.splitlines()),
                    ):
                        fallback_doc.add_paragraph(line)
                    buffer = BytesIO()
                    fallback_doc.save(buffer)
                    buffer.seek(0)
                    docx_bytes = buffer.getvalue()
                    buffer.close()
                    if docx_bytes:
                        docx_content = ContentFile(docx_bytes, name="tmp.docx")
                    else:
                        logger.error(
                            "El fallback de DOCX también produjo contenido vacío.",
                            extra={
                                "informe_pk": getattr(informe, "pk", None),
                                "tipo": tipo,
                            },
                        )
                except Exception as fallback_exc:
                    logger.error(
                        "No se pudo generar el fallback simplificado de DOCX.",
                        exc_info=fallback_exc,
                        extra={
                            "informe_pk": getattr(informe, "pk", None),
                            "tipo": tipo,
                        },
                    )

            base_filename = (
                slugify(f"{tipo}-informe-{informe.id}") or f"informe-{informe.id}"
            )
            pdf_content = ContentFile(pdf_bytes, name=f"{base_filename}.pdf")

            defaults = {
                "tipo": tipo,
                "informe_id": informe.id,
                "comedor": getattr(informe.admision, "comedor", None),
                "archivo": pdf_content,
            }

            if docx_content and any(
                f.name == "archivo_docx" for f in InformeTecnicoPDF._meta.get_fields()
            ):
                docx_content.name = f"{base_filename}.docx"
                defaults["archivo_docx"] = docx_content

            InformeTecnicoPDF.objects.update_or_create(
                admision=informe.admision,
                defaults=defaults,
            )

        except Exception:
            html_dump_path = None
            try:
                if "html_pdf" in locals() and html_pdf:
                    with tempfile.NamedTemporaryFile(
                        mode="w",
                        encoding="utf-8",
                        suffix=".html",
                        prefix="informe_tecnico_",
                        delete=False,
                    ) as temp_file:
                        temp_file.write(html_pdf)
                        html_dump_path = temp_file.name
            except Exception as dump_exc:
                logger.error(
                    "No se pudo escribir el HTML temporal para depuración: %s",
                    dump_exc,
                    extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
                )

            extra = {
                "informe_pk": getattr(informe, "pk", None),
                "tipo": tipo,
            }
            if html_dump_path:
                extra["html_dump_path"] = html_dump_path

            logger.exception("Error en generar_y_guardar_pdf", extra=extra)

    @staticmethod
    def get_informe_create_context(admision_id, tipo):
        try:
            admision = get_object_or_404(
                Admision.objects.select_related("comedor"), pk=admision_id
            )
            anexof = InformeService.get_anexo_form(admision)
            return {
                "tipo": tipo,
                "admision": admision,
                "comedor": admision.comedor,
                "anexof": anexof,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_create_context",
                extra={"admision_pk": admision_id, "tipo": tipo},
            )
            return {}

    @staticmethod
    def get_informe_update_context(informe, tipo):
        try:
            campos_a_subsanar = CampoASubsanar.objects.filter(
                informe=informe
            ).values_list("campo", flat=True)
            try:
                observacion = ObservacionGeneralInforme.objects.get(informe=informe)
            except ObservacionGeneralInforme.DoesNotExist:
                observacion = None

            anexof = InformeService.get_anexo_form(informe.admision)

            return {
                "tipo": tipo,
                "admision": informe.admision,
                "comedor": informe.admision.comedor,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "anexof": anexof,
                "campos_a_subsanar": list(campos_a_subsanar),
                "observacion": observacion,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_update_context",
                extra={
                    "informe_pk": getattr(informe, "pk", None),
                    "tipo": tipo,
                },
            )
            return {}

    @staticmethod
    @transaction.atomic
    def guardar_informe_y_anexo(
        form, admision, request_post, request_files=None, es_creacion=False, action=None
    ):
        """Guarda de forma atómica el informe técnico y su anexo asociado."""

        require_full = action == "submit"
        anexof = InformeService.get_anexo_form(
            admision, request_post, request_files, require_full=require_full
        )

        if require_full and not anexof.is_valid():
            return {"success": False, "anexof": anexof}

        if es_creacion:
            InformeService.preparar_informe_para_creacion(
                form.instance, admision.id, action
            )
        else:
            InformeService.verificar_estado_para_revision(form.instance, action)

        informe = form.save(commit=False)
        informe.admision = admision
        informe.save()
        if hasattr(form, "save_m2m"):
            form.save_m2m()

        if anexof.is_valid():
            anexo = anexof.save(commit=False)
        else:
            anexo = anexof.instance or Anexo()

        anexo.admision = admision
        anexo.save()
        if hasattr(anexof, "save_m2m") and anexof.is_valid():
            anexof.save_m2m()

        return {"success": True, "informe": informe, "anexof": anexof, "anexo": anexo}

    @staticmethod
    def get_context_informe_detail(informe, tipo):
        try:
            return {
                "tipo": tipo,
                "admision": informe.admision,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "pdf": InformeTecnicoPDF.objects.filter(
                    admision=informe.admision, tipo=tipo, informe_id=informe.id
                ).first(),
            }
        except Exception:
            logger.exception(
                "Error en get_context_informe_detail",
                extra={
                    "informe_pk": getattr(informe, "pk", None),
                    "tipo": tipo,
                },
            )
            return {}

    @staticmethod
    def procesar_revision_informe(request, tipo, informe):
        try:
            nuevo_estado = request.POST.get("estado")
            if nuevo_estado not in ["A subsanar", "Validado"]:
                return

            InformeService.actualizar_estado_informe(informe, nuevo_estado, tipo)

            if nuevo_estado == "A subsanar":
                campos_a_subsanar = request.POST.getlist("campos_a_subsanar")
                observacion = request.POST.get("observacion", "").strip()

                CampoASubsanar.objects.filter(informe=informe).delete()

                for campo in campos_a_subsanar:
                    CampoASubsanar.objects.create(informe=informe, campo=campo)

                obs_obj, _ = ObservacionGeneralInforme.objects.get_or_create(
                    informe=informe
                )
                obs_obj.texto = observacion
                obs_obj.save()
            else:
                CampoASubsanar.objects.filter(informe=informe).delete()
                ObservacionGeneralInforme.objects.filter(informe=informe).delete()
        except Exception:
            logger.exception(
                "Error en procesar_revision_informe",
                extra={
                    "informe_pk": getattr(informe, "pk", None),
                    "tipo": tipo,
                },
            )

    @staticmethod
    def guardar_campos_complementarios(informe_tecnico, campos_dict, usuario):
        try:
            informe = InformeComplementario.objects.create(
                admision=informe_tecnico.admision,
                informe_tecnico=informe_tecnico,
                creado_por=usuario,
            )
            for campo, valor in campos_dict.items():
                InformeComplementarioCampos.objects.create(
                    campo=campo, value=valor, informe_complementario=informe
                )
            return informe
        except Exception:
            logger.exception(
                "Error en guardar_campos_complementarios",
                extra={"informe_tecnico_pk": getattr(informe_tecnico, "pk", None)},
            )
            return None

    @staticmethod
    def generar_y_guardar_pdf_complementario(informe_complementario):
        try:
            campos = [
                (campo.campo, campo.value)
                for campo in informe_complementario.informecomplementariocampos_set.all()
            ]
            html_string = render_to_string(
                "admisiones2/pdf_informe_complementario.html",
                {"informe": informe_complementario, "campos": campos},
            )
            pdf_bytes = HTML(string=html_string).write_pdf()
            nombre_archivo = f"complementario_{informe_complementario.id}.pdf"
            informe_complementario.pdf.save(
                nombre_archivo, ContentFile(pdf_bytes), save=True
            )
        except Exception:
            logger.exception(
                "Error en generar_y_guardar_pdf_complementario",
                extra={
                    "informe_complementario_pk": getattr(
                        informe_complementario, "pk", None
                    )
                },
            )
