from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string, get_template
from datetime import date, datetime

from django.utils import timezone
from django.utils.text import slugify
from django.utils.html import strip_tags
from django.conf import settings
from weasyprint import HTML
from django.core.files.base import ContentFile
from io import BytesIO
from docx import Document
from htmldocx import HtmlToDocx
from .docx_service import DocumentTemplateService, AdmisionesContextService
from django.template import TemplateDoesNotExist
from django.db import transaction
from django.contrib import messages
import logging
import tempfile
import traceback
import os

from ..utils import generar_texto_comidas

logger = logging.getLogger("admisiones.services.informes")

from admisiones.models.admisiones import (
    InformeTecnico,
    CampoASubsanar,
    ObservacionGeneralInforme,
    InformeTecnicoPDF,
    Admision,
    InformeComplementario,
    InformeComplementarioCampos,
    InformeTecnicoComplementarioPDF,
)
from admisiones.forms.admisiones_forms import (
    InformeTecnicoJuridicoForm,
    InformeTecnicoBaseForm,
)


class InformeService:
    @staticmethod
    def _get_base_url():
        """Helper to get base URL for PDF generation"""
        return str(
            getattr(settings, "STATIC_ROOT", "")
            or getattr(settings, "BASE_DIR", "")
            or "."
        )

    @staticmethod
    def _generate_docx_content(html_content, informe_pk=None):
        """Helper to generate DOCX content with fallback"""
        try:
            doc = Document()
            HtmlToDocx().add_html_to_document(html_content, doc)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.getvalue()
            buffer.close()
            if docx_bytes:
                return ContentFile(docx_bytes, name="tmp.docx")
        except Exception as e:
            logger.warning(
                f"Error generando DOCX: {str(e)}", extra={"informe_pk": informe_pk}
            )
            try:
                fallback_doc = Document()
                fallback_text = strip_tags(html_content)
                for line in filter(
                    None, (segment.strip() for segment in fallback_text.splitlines())
                ):
                    fallback_doc.add_paragraph(line)
                buffer = BytesIO()
                fallback_doc.save(buffer)
                buffer.seek(0)
                docx_bytes = buffer.getvalue()
                buffer.close()
                if docx_bytes:
                    return ContentFile(docx_bytes, name="tmp.docx")
            except Exception:
                logger.error(
                    "Fallback DOCX generation failed", extra={"informe_pk": informe_pk}
                )
        return None

    @staticmethod
    def get_form_class_por_tipo(tipo):
        return (
            InformeTecnicoJuridicoForm if tipo == "juridico" else InformeTecnicoBaseForm
        )

    @staticmethod
    def get_tipo_from_kwargs(kwargs):
        return kwargs.get("tipo", "base")

    @staticmethod
    def get_queryset_informe_por_tipo(tipo):
        return InformeTecnico.objects.filter(tipo=tipo)

    @staticmethod
    def get_admision_y_tipo_from_kwargs(kwargs):
        try:
            tipo = kwargs.get("tipo", "base")
            admision_id = kwargs.get("admision_id")
            admision = get_object_or_404(Admision, pk=admision_id)
            return admision, tipo
        except Exception:
            logger.exception(
                "Error en get_admision_y_tipo_from_kwargs", extra={"kwargs": kwargs}
            )
            return None, "base"

    @staticmethod
    def verificar_estado_para_revision(informe, action=None):
        """Actualiza estado al modificar un informe existente según la acción."""
        try:
            if action == "draft":
                informe.estado_formulario = "borrador"
                informe.estado = "Iniciado"
            else:
                if informe.estado != "Validado":
                    CampoASubsanar.objects.filter(informe=informe).delete()
                    ObservacionGeneralInforme.objects.filter(informe=informe).delete()
                    informe.estado_formulario = "finalizado"
                    informe.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en verificar_estado_para_revision",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )

    @staticmethod
    def get_campos_visibles_informe(informe):
        try:
            campos_excluidos_comunes = [
                "id",
                "admision",
                "estado",
                "tipo",
                "estado_formulario",
            ]

            if informe.tipo == "juridico":
                campos_excluidos_especificos = [
                    "declaracion_jurada_recepcion_subsidios",
                    "constancia_inexistencia_percepcion_otros_subsidios",
                    "organizacion_avalista_1",
                    "organizacion_avalista_2",
                    "material_difusion_vinculado",
                    "if_relevamiento",
                ]
            elif informe.tipo == "base":
                campos_excluidos_especificos = [
                    "validacion_registro_nacional",
                    "IF_relevamiento_territorial",
                ]
            else:
                campos_excluidos_especificos = []

            campos_excluidos = campos_excluidos_comunes + campos_excluidos_especificos

            return [
                (
                    field.verbose_name,
                    InformeService._formatear_valor_campo(informe, field),
                )
                for field in informe._meta.fields
                if field.name not in campos_excluidos
            ]
        except Exception:
            logger.exception(
                "Error en get_campos_visibles_informe",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return []

    @staticmethod
    def _formatear_valor_campo(informe, field):
        value = getattr(informe, field.name)

        if field.choices:
            display_value = getattr(informe, f"get_{field.name}_display", None)
            if display_value:
                try:
                    return display_value()
                except Exception:
                    pass

        if isinstance(value, bool):
            return "Sí" if value else "No"

        if isinstance(value, datetime):
            if timezone.is_aware(value):
                value = timezone.localtime(value)
            return value.strftime("%d/%m/%Y %H:%M")

        if isinstance(value, date):
            return value.strftime("%d/%m/%Y")

        return value

    @staticmethod
    def preparar_informe_para_creacion(instance, admision_id, action=None):
        """Inicializa un informe técnico nuevo según la acción (borrador/finalizado)."""
        try:
            instance.admision_id = admision_id
            if action == "draft":
                instance.estado_formulario = "borrador"
                instance.estado = "Iniciado"
            else:
                instance.estado_formulario = "finalizado"
                instance.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en preparar_informe_para_creacion",
                extra={"admision_pk": admision_id},
            )

    @staticmethod
    def get_informe_por_tipo_y_pk(tipo, pk):
        try:
            return get_object_or_404(InformeTecnico, tipo=tipo, pk=pk)
        except Exception:
            logger.exception(
                "Error en get_informe_por_tipo_y_pk",
                extra={"tipo": tipo, "informe_pk": pk},
            )
            return None

    @staticmethod
    def actualizar_estado_informe(informe, nuevo_estado, tipo=None):
        try:
            informe.estado = nuevo_estado
            informe.save()

            # Actualizar estado de admisión según el nuevo estado del informe
            from .admisiones_service import AdmisionService

            if nuevo_estado == "A subsanar":
                AdmisionService.actualizar_estado_admision(
                    informe.admision, "subsanar_informe"
                )
            elif nuevo_estado == "Validado":
                AdmisionService.actualizar_estado_admision(
                    informe.admision, "aprobar_informe_tecnico"
                )
                # No generar nuevos archivos - el DOCX del técnico es el final
        except Exception as e:
            logger.exception(
                f"Error en actualizar_estado_informe: {str(e)}",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            raise

    @staticmethod
    def _normalizar_tipo_admision(admision):
        tipo = getattr(admision, "tipo", None)
        normalizado = (tipo or "").strip().lower()
        return normalizado or "incorporacion"

    @staticmethod
    def generar_docx_con_template(informe, template_name=None):
        """Genera DOCX usando template con docxtpl"""
        try:
            # Seleccionar template basado en tipo de admisión e informe
            if not template_name:
                admision_tipo = InformeService._normalizar_tipo_admision(
                    informe.admision
                )
                informe_tipo = informe.tipo
                template_name = (
                    f"{admision_tipo}_docx_informe_tecnico_{informe_tipo}.docx"
                )

            context = AdmisionesContextService.preparar_contexto_informe_tecnico(
                informe
            )
            logger.debug(
                "Generando DOCX para informe %s con template %s (tipo %s)",
                informe.id,
                template_name,
                informe.tipo,
            )

            return DocumentTemplateService.generar_docx(template_name, context)
        except Exception as e:
            logger.exception(
                "No se pudo procesar template DOCX para informe %s: %s",
                getattr(informe, "pk", None),
                str(e),
            )
            return None

    @staticmethod
    def generar_y_guardar_pdf(informe, tipo):
        """
        Genera y guarda PDF y DOCX del informe técnico
        """
        try:
            context = {
                "informe": informe,
                "texto_comidas": generar_texto_comidas(informe),
            }

            # Mapear tipos de informe para templates
            admision_tipo = InformeService._normalizar_tipo_admision(informe.admision)
            informe_tipo_map = {
                "base": "base",
                "juridico": "juridico",
                "juridico eclesiastico": "juridico",
            }
            informe_tipo = informe_tipo_map.get(informe.tipo, "base")

            pdf_template = f"admisiones/pdf/{admision_tipo}_pdf_informe_tecnico_{informe_tipo}.html"
            docx_template = f"admisiones/docx/{admision_tipo}_docx_informe_tecnico_{informe_tipo}.html"

            logger.debug(
                "Generando archivos para informe %s (PDF: %s, DOCX: %s)",
                informe.id,
                pdf_template,
                docx_template,
            )

            # Generar PDF
            try:
                html_pdf = render_to_string(pdf_template, context)
                if not html_pdf.strip():
                    raise ValueError(f"Template PDF vacío: {pdf_template}")

                pdf_bytes = HTML(
                    string=html_pdf, base_url=InformeService._get_base_url()
                ).write_pdf()
                if not pdf_bytes:
                    raise ValueError("WeasyPrint no generó contenido PDF")

            except Exception as e:
                logger.error("Error generando PDF: %s", str(e))
                raise

            # Generar DOCX HTML template
            try:
                html_docx = render_to_string(docx_template, context)
                if not html_docx.strip():
                    html_docx = html_pdf
            except Exception as e:
                logger.warning("Error con template DOCX HTML: %s", str(e))
                html_docx = html_pdf

            # Generar DOCX con docxtpl
            docx_content = None
            try:
                logger.info(
                    f"Intentando generar DOCX con template para informe {informe.id}"
                )
                docx_buffer = InformeService.generar_docx_con_template(informe)
                if docx_buffer:
                    logger.info("DOCX generado exitosamente con template docxtpl")
                    docx_content = ContentFile(docx_buffer.getvalue(), name="tmp.docx")
                else:
                    raise ValueError("Template DOCX retornó None")
            except Exception as e:
                logger.warning("Template DOCX falló: %s, usando fallback HTML", str(e))
                docx_content = InformeService._generate_docx_content(
                    html_docx, getattr(informe, "pk", None)
                )

            # Guardar archivos - PDF final (sin prefijo "borrador")
            base_filename = (
                slugify(f"{tipo}-informe-{informe.id}") or f"informe-{informe.id}"
            )
            pdf_content = ContentFile(pdf_bytes, name=f"{base_filename}.pdf")

            defaults = {
                "tipo": tipo,
                "informe_id": informe.id,
                "comedor": getattr(informe.admision, "comedor", None),
                "archivo": pdf_content,
            }

            if docx_content:
                docx_content.name = f"{base_filename}.docx"
                defaults["archivo_docx"] = docx_content
                logger.debug(
                    "Archivos PDF y DOCX preparados para guardar para informe %s",
                    informe.id,
                )
            else:
                logger.debug(
                    "Solo se guardará PDF para informe %s, generación DOCX falló",
                    informe.id,
                )

            pdf_obj, created = InformeTecnicoPDF.objects.update_or_create(
                admision=informe.admision, defaults=defaults
            )

            action = "creado" if created else "actualizado"
            logger.info(
                "InformeTecnicoPDF %s exitosamente para informe %s",
                action,
                informe.id,
            )

            return pdf_obj

        except Exception as e:
            logger.exception(
                f"Error crítico en generar_y_guardar_pdf: {str(e)}",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            raise

    @staticmethod
    def get_informe_create_context(admision_id, tipo):
        try:
            admision = get_object_or_404(
                Admision.objects.select_related("comedor"), pk=admision_id
            )
            return {
                "tipo": tipo,
                "admision": admision,
                "comedor": admision.comedor,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_create_context",
                extra={"admision_pk": admision_id, "tipo": tipo},
            )
            return {}

    @staticmethod
    def get_informe_update_context(informe, tipo):
        try:
            campos_a_subsanar_db = CampoASubsanar.objects.filter(
                informe=informe
            ).values_list("campo", flat=True)

            # Incluir tanto nombres de campo como verbose names para el template
            campos_a_subsanar = list(campos_a_subsanar_db)
            field_to_verbose = {
                field.name: field.verbose_name for field in informe._meta.fields
            }

            # Agregar también los verbose names para mostrar en la interfaz
            for campo in campos_a_subsanar_db:
                verbose_name = field_to_verbose.get(campo, campo)
                if verbose_name not in campos_a_subsanar:
                    campos_a_subsanar.append(verbose_name)

            try:
                observacion = ObservacionGeneralInforme.objects.get(informe=informe)
            except ObservacionGeneralInforme.DoesNotExist:
                observacion = None

            return {
                "tipo": tipo,
                "admision": informe.admision,
                "comedor": informe.admision.comedor,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "campos_a_subsanar": campos_a_subsanar,
                "observacion": observacion,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_update_context",
                extra={
                    "informe_pk": getattr(informe, "pk", None),
                    "tipo": tipo,
                },
            )
            return {}

    @staticmethod
    @transaction.atomic
    def guardar_informe(form, admision, es_creacion=False, action=None, usuario=None):
        """Guarda el informe técnico."""
        try:
            if es_creacion:
                existente = (
                    InformeTecnico.objects.filter(
                        admision=admision,
                        tipo=form.instance.tipo,
                    )
                    .order_by("-id")
                    .first()
                )
                if existente:
                    form.instance.pk = existente.pk
                    form.instance._state.adding = False
                    es_creacion = False

            if es_creacion:
                InformeService.preparar_informe_para_creacion(
                    form.instance, admision.id, action
                )
            else:
                InformeService.verificar_estado_para_revision(form.instance, action)

            informe = form.save(commit=False)
            informe.admision = admision

            # Set creado_por and modificado_por fields
            if usuario:
                if es_creacion:
                    informe.creado_por = usuario
                informe.modificado_por = usuario

            # Limpiar observaciones de subsanación cuando se finaliza el informe
            if action == "submit" and informe.observaciones_subsanacion:
                informe.observaciones_subsanacion = None
                
            informe.save()
            if hasattr(form, "save_m2m"):
                form.save_m2m()

            if action == "submit" and informe.estado_formulario == "finalizado":
                resultado_docx = InformeService.generar_docx_borrador(informe)
                if resultado_docx:
                    # Solo actualizar estado si el DOCX se generó exitosamente
                    admision.estado_admision = "informe_tecnico_finalizado"
                    admision.save()
                else:
                    raise Exception("No se pudo generar el DOCX borrador")

            # Actualizar estado de admisión según la acción
            from .admisiones_service import AdmisionService

            if es_creacion and action != "submit":
                AdmisionService.actualizar_estado_admision(
                    admision, "iniciar_informe_tecnico"
                )

            return {"success": True, "informe": informe}
        except Exception as e:
            logger.exception(
                "Error en guardar_informe",
                extra={"admision_pk": getattr(admision, "pk", None)},
            )
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_context_informe_detail(informe, tipo):
        try:
            pdf_filter = {
                "admision": informe.admision,
                "tipo": tipo,
                "informe_id": informe.id,
            }
            pdf_final = (
                InformeTecnicoPDF.objects.filter(**pdf_filter).first()
                if informe.estado == "Validado"
                else None
            )
            pdf_borrador = (
                InformeTecnicoPDF.objects.filter(**pdf_filter).first()
                if informe.estado_formulario == "finalizado"
                and informe.estado != "Validado"
                else None
            )

            return {
                "tipo": tipo,
                "admision": informe.admision,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "pdf": pdf_final,
                "pdf_borrador": pdf_borrador,
            }
        except Exception:
            logger.exception(
                "Error en get_context_informe_detail",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            return {}

    @staticmethod
    def procesar_revision_informe(request, tipo, informe):
        try:
            nuevo_estado = request.POST.get("estado")
            if nuevo_estado not in ["A subsanar", "Validado"]:
                return

            InformeService.actualizar_estado_informe(informe, nuevo_estado, tipo)

            if nuevo_estado == "A subsanar":
                campos_a_subsanar = request.POST.getlist("campos_a_subsanar")
                observacion = request.POST.get("observacion", "").strip()

                CampoASubsanar.objects.filter(informe=informe).delete()

                verbose_to_field = {
                    field.verbose_name: field.name for field in informe._meta.fields
                }
                for campo in campos_a_subsanar:
                    field_name = verbose_to_field.get(campo, campo)
                    CampoASubsanar.objects.create(informe=informe, campo=field_name)

                # Guardar observaciones de subsanación en el informe
                informe.observaciones_subsanacion = observacion
                informe.save()
                
                obs_obj, _ = ObservacionGeneralInforme.objects.get_or_create(
                    informe=informe
                )
                obs_obj.texto = observacion
                obs_obj.save()
            else:
                CampoASubsanar.objects.filter(informe=informe).delete()
                ObservacionGeneralInforme.objects.filter(informe=informe).delete()
                # Limpiar observaciones de subsanación
                informe.observaciones_subsanacion = None
                informe.save()
        except Exception:
            logger.exception(
                "Error en procesar_revision_informe",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )

    @staticmethod
    def guardar_campos_complementarios(informe_tecnico, campos_dict, usuario):
        """
        Guarda los campos modificados como un solo conjunto de cambios.
        Siempre actualiza el informe complementario existente o crea uno nuevo.
        """
        try:
            informe, created = InformeComplementario.objects.get_or_create(
                admision=informe_tecnico.admision,
                defaults={
                    "informe_tecnico": informe_tecnico,
                    "creado_por": usuario,
                    "estado": "borrador",
                },
            )

            InformeComplementarioCampos.objects.filter(
                informe_complementario=informe
            ).delete()

            for campo, valor in campos_dict.items():
                InformeComplementarioCampos.objects.create(
                    campo=campo, value=valor, informe_complementario=informe
                )

            return informe

        except Exception:
            logger.exception(
                "Error en guardar_campos_complementarios",
                extra={"informe_tecnico_pk": getattr(informe_tecnico, "pk", None)},
            )
            return None

    @staticmethod
    def generar_y_guardar_pdf_complementario(informe_complementario):
        try:
            with transaction.atomic():
                informe = informe_complementario.informe_tecnico
                campos_modificados = InformeComplementarioCampos.objects.filter(
                    informe_complementario=informe_complementario
                )

                verbose_to_field = {
                    field.verbose_name.lower().strip(): field.name
                    for field in informe._meta.fields
                }

                campos_actualizados_detalle = []

                for c in campos_modificados:
                    key = c.campo.lower().strip()
                    field_name = (
                        c.campo
                        if hasattr(informe, c.campo)
                        else verbose_to_field.get(key)
                    )
                    if not field_name:
                        continue

                    field = informe._meta.get_field(field_name)
                    valor_original = getattr(informe, field_name, None)
                    nuevo_valor = c.value

                    if field.get_internal_type() in [
                        "IntegerField",
                        "PositiveIntegerField",
                    ]:
                        try:
                            nuevo_valor = int(nuevo_valor) if nuevo_valor else 0
                        except ValueError:
                            nuevo_valor = None
                    elif field.get_internal_type() == "DateField":
                        from django.utils.dateparse import parse_date

                        parsed = parse_date(str(nuevo_valor))
                        if parsed:
                            nuevo_valor = parsed

                    setattr(informe, field_name, nuevo_valor)
                    campos_actualizados_detalle.append(
                        {
                            "campo": field.verbose_name or field.name,
                            "valor_anterior": valor_original,
                            "valor_nuevo": nuevo_valor,
                        }
                    )

                if campos_actualizados_detalle:
                    informe.save()

                usuario = informe_complementario.creado_por
                creado_por = usuario.username if usuario else "—"

                context = {
                    "informe": informe,
                    "texto_comidas": generar_texto_comidas(informe),
                    "campos_actualizados": campos_actualizados_detalle,
                    "fecha": getattr(informe_complementario, "creado", None),
                    "creado_por": creado_por,
                }

                pdf_template = "admisiones/pdf/informe_tecnico_complementario.html"
                html_pdf = render_to_string(pdf_template, context)
                pdf_bytes = HTML(
                    string=html_pdf, base_url=InformeService._get_base_url()
                ).write_pdf()

                docx_content = DocumentTemplateService.generar_docx(
                    template_name="informe_tecnico_complementario.docx",
                    context=context,
                    app_name="admisiones",
                )

                base_filename = slugify(
                    f"informe-complementario-{informe.id}-admision-{informe.admision_id}-{informe.tipo}"
                )

                defaults = {
                    "admision": informe.admision,
                    "tipo": informe.tipo,
                    "archivo": ContentFile(pdf_bytes, name=f"{base_filename}.pdf"),
                }

                if docx_content:
                    docx_content.seek(0)
                    defaults["archivo_docx"] = ContentFile(
                        docx_content.read(), name=f"{base_filename}.docx"
                    )

                InformeTecnicoComplementarioPDF.objects.update_or_create(
                    informe_complementario=informe_complementario,
                    defaults=defaults,
                )

                return True

        except Exception:
            import logging

            logging.exception("Error en generar_y_guardar_pdf_complementario")
            return None

    @staticmethod
    def generar_docx_borrador(informe):
        """Genera un DOCX borrador del informe técnico para edición del técnico"""
        try:
            # Generar DOCX con docxtpl
            docx_content = None
            try:
                docx_buffer = InformeService.generar_docx_con_template(informe)
                if docx_buffer:
                    docx_content = ContentFile(docx_buffer.getvalue(), name="tmp.docx")
                else:
                    raise ValueError("Template DOCX retornó None")
            except Exception as e:
                logger.warning("Template DOCX falló: %s, usando fallback HTML", str(e))
                # Fallback: generar desde HTML
                context = {
                    "informe": informe,
                    "texto_comidas": generar_texto_comidas(informe),
                }
                
                admision_tipo = InformeService._normalizar_tipo_admision(informe.admision)
                informe_tipo_map = {
                    "base": "base",
                    "juridico": "juridico",
                    "juridico eclesiastico": "juridico",
                }
                informe_tipo = informe_tipo_map.get(informe.tipo, "base")
                docx_template = f"admisiones/docx/{admision_tipo}_docx_informe_tecnico_{informe_tipo}.html"
                
                try:
                    html_docx = render_to_string(docx_template, context)
                    docx_content = InformeService._generate_docx_content(
                        html_docx, getattr(informe, "pk", None)
                    )
                except Exception:
                    logger.error("Fallback DOCX generation también falló")
                    return None

            if not docx_content:
                return None

            # Guardar DOCX borrador
            base_filename = slugify(
                f"informe-{informe.id}-admision-{informe.admision_id}-{informe.tipo}-borrador"
            )
            docx_content.name = f"{base_filename}.docx"

            # Solo crear/actualizar si no existe un PDF final (validado)
            if informe.estado != "Validado":
                pdf_borrador, created = InformeTecnicoPDF.objects.update_or_create(
                    admision=informe.admision,
                    defaults={
                        "tipo": informe.tipo,
                        "informe_id": informe.id,
                        "comedor": informe.admision.comedor,
                        "archivo_docx": docx_content,
                    },
                )
                
                # Actualizar estado del informe
                informe.estado = "Docx generado"
                informe.save()
                
                logger.info(
                    "DOCX borrador generado para informe %s, estado actualizado",
                    informe.id,
                )
                
                return pdf_borrador
            else:
                logger.info(
                    "No se genera DOCX borrador porque ya existe PDF final validado para informe %s",
                    informe.id,
                )
                return None

        except Exception:
            logger.exception(
                "Error en generar_docx_borrador",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return None

    @staticmethod
    def subir_docx_editado(informe, archivo_docx, usuario=None):
        """Maneja la subida del DOCX editado por el técnico"""
        try:
            # Obtener o crear el registro PDF
            pdf_obj, created = InformeTecnicoPDF.objects.get_or_create(
                admision=informe.admision,
                defaults={
                    "tipo": informe.tipo,
                    "informe_id": informe.id,
                    "comedor": informe.admision.comedor,
                }
            )
            
            # Pisar el DOCX borrador con el editado
            base_filename = slugify(
                f"informe-{informe.id}-admision-{informe.admision_id}-{informe.tipo}-final"
            )
            archivo_docx.name = f"{base_filename}.docx"
            pdf_obj.archivo_docx = archivo_docx
            pdf_obj.save()
            
            # Actualizar estado del informe
            informe.estado = "Docx editado"
            informe.save()
            
            # Actualizar estado de admisión - ahora pasa a revisión
            from .admisiones_service import AdmisionService
            AdmisionService.actualizar_estado_admision(
                informe.admision, "enviar_informe_revision"
            )
            
            logger.info(
                "DOCX editado subido exitosamente para informe %s",
                informe.id,
            )
            
            return pdf_obj
            
        except Exception:
            logger.exception(
                "Error en subir_docx_editado",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return None

    @staticmethod
    def obtener_cambios_complementarios_texto(informe_complementario):
        """Obtiene los cambios del complementario como texto para mostrar en la interfaz"""
        try:
            return list(
                InformeComplementarioCampos.objects.filter(
                    informe_complementario=informe_complementario
                )
            )
        except Exception:
            logger.exception(
                "Error en obtener_cambios_complementarios_texto",
                extra={
                    "informe_complementario_pk": getattr(
                        informe_complementario, "pk", None
                    )
                },
            )
            return []
