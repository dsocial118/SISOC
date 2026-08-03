import logging
import os
import traceback
from datetime import date, datetime
from io import BytesIO
from typing import Any

from django.conf import settings
from django.contrib import messages
from django.core.files.base import ContentFile
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.template import Context, Engine, TemplateSyntaxError
from django.template.loader import get_template, render_to_string
from django.utils import timezone
from django.utils.html import strip_tags
from django.utils.text import slugify
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, Pt
from htmldocx import HtmlToDocx

from ...utils import generar_texto_comidas
from ..docx_service import AdmisionesContextService, DocumentTemplateService
from admisiones.forms.admisiones_forms import (
    InformeTecnicoBaseForm,
    InformeTecnicoJuridicoForm,
)
from admisiones.models.admisiones import (
    Admision,
    CampoASubsanar,
    InformeComplementario,
    InformeComplementarioCampos,
    InformeTecnico,
    InformeTecnicoComplementarioPDF,
    InformeTecnicoPDF,
    ObservacionGeneralInforme,
)

logger = logging.getLogger("admisiones.services.informes")

MATRICES_DE_PRESTACIONES = {
    "Prestaciones aprobadas en el último convenio": (
        "prestaciones_ultimo_convenio",
        "aprobadas_ultimo_convenio",
    ),
    "Solicitudes": ("solicitudes", "solicitudes"),
    "Prestaciones Aprobadas": ("prestaciones_aprobadas", "aprobadas"),
}
COMIDAS = ("desayuno", "almuerzo", "merienda", "cena")
DIAS_SEMANA = (
    "lunes",
    "martes",
    "miercoles",
    "jueves",
    "viernes",
    "sabado",
    "domingo",
)


class InformeService:
    COLOR_BORDE_TABLA_TEMPLATE = "7A8EA1"
    COLOR_ENCABEZADO_TABLA_TEMPLATE = "EAF3F2"
    ANCHO_PAGINA_TEMPLATE_MM = 210
    ALTO_PAGINA_TEMPLATE_MM = 297
    MARGEN_PAGINA_TEMPLATE_MM = 20
    FUENTE_BASE_TEMPLATE = "Times New Roman"
    TAMANIO_FUENTE_BASE_TEMPLATE_PT = 12

    @staticmethod
    def _get_base_url():
        """Helper to get base URL for PDF generation"""
        return str(
            getattr(settings, "STATIC_ROOT", "")
            or getattr(settings, "BASE_DIR", "")
            or "."
        )

    @staticmethod
    def _generate_docx_content(
        html_content,
        informe_pk=None,
        estilizar_tablas_templates=False,
    ):
        """Helper to generate DOCX content with fallback"""
        try:
            doc = Document()
            if estilizar_tablas_templates:
                InformeService._configurar_documento_template(doc)
            HtmlToDocx().add_html_to_document(html_content, doc)
            if estilizar_tablas_templates:
                InformeService._aplicar_estilo_tablas_templates(doc)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.getvalue()
            buffer.close()
            if docx_bytes:
                return ContentFile(docx_bytes, name="tmp.docx")
        except Exception as e:
            logger.warning(
                f"Error generando DOCX: {str(e)}", extra={"informe_pk": informe_pk}
            )
            try:
                fallback_doc = Document()
                if estilizar_tablas_templates:
                    InformeService._configurar_documento_template(fallback_doc)
                fallback_text = strip_tags(html_content)
                for line in filter(
                    None, (segment.strip() for segment in fallback_text.splitlines())
                ):
                    fallback_doc.add_paragraph(line)
                buffer = BytesIO()
                fallback_doc.save(buffer)
                buffer.seek(0)
                docx_bytes = buffer.getvalue()
                buffer.close()
                if docx_bytes:
                    return ContentFile(docx_bytes, name="tmp.docx")
            except Exception:
                logger.error(
                    "Fallback DOCX generation failed", extra={"informe_pk": informe_pk}
                )
        return None

    @staticmethod
    def get_form_class_por_tipo(tipo):
        return (
            InformeTecnicoJuridicoForm if tipo == "juridico" else InformeTecnicoBaseForm
        )

    @staticmethod
    def get_tipo_from_kwargs(kwargs):
        return kwargs.get("tipo", "base")

    @staticmethod
    def get_queryset_informe_por_tipo(tipo):
        return InformeTecnico.objects.filter(tipo=tipo)

    @staticmethod
    def get_admision_y_tipo_from_kwargs(kwargs):
        try:
            tipo = kwargs.get("tipo", "base")
            admision_id = kwargs.get("admision_id")
            admision = get_object_or_404(Admision, pk=admision_id)
            return admision, tipo
        except Exception:
            logger.exception(
                "Error en get_admision_y_tipo_from_kwargs", extra={"kwargs": kwargs}
            )
            return None, "base"

    @staticmethod
    def verificar_estado_para_revision(informe, action=None):
        """Actualiza estado al modificar un informe existente según la acción."""
        try:
            if action == "draft":
                informe.estado_formulario = "borrador"
                informe.estado = "Iniciado"
            else:
                if informe.estado != "Validado":
                    CampoASubsanar.objects.filter(informe=informe).delete()
                    ObservacionGeneralInforme.objects.filter(informe=informe).delete()
                    informe.estado_formulario = "finalizado"
                    informe.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en verificar_estado_para_revision",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )

    @staticmethod
    def get_campos_visibles_informe(informe):
        try:
            campos_excluidos_comunes = [
                "id",
                "admision",
                "estado",
                "tipo",
                "estado_formulario",
            ]

            if informe.tipo == "juridico":
                campos_excluidos_especificos = [
                    "declaracion_jurada_recepcion_subsidios",
                    "constancia_inexistencia_percepcion_otros_subsidios",
                    "organizacion_avalista_1",
                    "organizacion_avalista_2",
                    "material_difusion_vinculado",
                    "if_relevamiento",
                ]
            elif informe.tipo == "base":
                campos_excluidos_especificos = [
                    "validacion_registro_nacional",
                    "IF_relevamiento_territorial",
                ]
            else:
                campos_excluidos_especificos = []

            campos_excluidos = campos_excluidos_comunes + campos_excluidos_especificos

            return [
                (
                    field.verbose_name,
                    InformeService._formatear_valor_campo(informe, field),
                )
                for field in informe._meta.fields
                if field.name not in campos_excluidos
            ]
        except Exception:
            logger.exception(
                "Error en get_campos_visibles_informe",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return []

    @staticmethod
    def get_campos_agrupados_informe(informe) -> list[dict[str, Any]]:
        """Agrupa campos visibles y conserva matrices de prestaciones por día."""
        titulo_campos_especificos = {
            "base": "Campos Específicos - Organización de Base",
            "juridico": "Campos Específicos - Organización Jurídica",
        }.get(getattr(informe, "tipo", None), "Campos Específicos")
        grupos = {
            "Datos de la Organización": [],
            "Datos del Representante": [],
            "Datos del Comedor/Merendero": [],
            "Responsable de la Tarjeta": [],
            "Prestaciones aprobadas en el último convenio": [],
            "Solicitudes": [],
            "Prestaciones Aprobadas": [],
            "Información Adicional": [],
            titulo_campos_especificos: [],
            "Resolución de pago": [],
        }
        valores = dict(InformeService.get_campos_visibles_informe(informe))
        es_renovacion = getattr(getattr(informe, "admision", None), "tipo", None) == (
            "renovacion"
        )
        campos_especificos = {
            "declaracion_jurada_recepcion_subsidios",
            "constancia_inexistencia_percepcion_otros_subsidios",
            "organizacion_avalista_1",
            "organizacion_avalista_2",
            "material_difusion_vinculado",
            "if_relevamiento",
            "validacion_registro_nacional",
            "IF_relevamiento_territorial",
        }
        for field in informe._meta.fields:
            nombre = str(field.verbose_name)
            if nombre not in valores:
                continue
            es_campo_renovacion = field.name.startswith(
                ("aprobadas_ultimo_convenio_", "resolucion_de_pago_", "monto_")
            )
            if es_campo_renovacion and not es_renovacion:
                continue
            if field.name in campos_especificos:
                grupo = titulo_campos_especificos
            elif field.name == "expediente_nro" or "organizacion" in field.name:
                grupo = "Datos de la Organización"
            elif field.name.startswith("representante_"):
                grupo = "Datos del Representante"
            elif field.name.startswith("responsable_tarjeta_"):
                grupo = "Responsable de la Tarjeta"
            elif field.name.startswith("aprobadas_ultimo_convenio_"):
                grupo = "Prestaciones aprobadas en el último convenio"
            elif field.name.startswith("solicitudes_"):
                grupo = "Solicitudes"
            elif field.name.startswith("aprobadas_"):
                grupo = "Prestaciones Aprobadas"
            elif field.name.startswith(("resolucion_de_pago_", "monto_")):
                grupo = "Resolución de pago"
            elif field.name.endswith("_espacio") or field.name in {
                "tipo_espacio",
                "nombre_espacio",
                "barrio_espacio",
            }:
                grupo = "Datos del Comedor/Merendero"
            else:
                grupo = "Información Adicional"
            grupos[grupo].append(
                {
                    "identificador": field.name,
                    "nombre": nombre,
                    "valor": valores[nombre],
                }
            )

        secciones = []
        for titulo, campos in grupos.items():
            if not campos:
                continue
            matriz = MATRICES_DE_PRESTACIONES.get(titulo)
            if not matriz:
                secciones.append(
                    {
                        "titulo": titulo,
                        "tipo": "campos",
                        "campos": campos,
                    }
                )
                continue

            identificador, prefijo = matriz
            campos_por_identificador = {
                campo["identificador"]: campo for campo in campos
            }
            filas = []
            for comida in COMIDAS:
                campos_fila = [
                    campos_por_identificador[f"{prefijo}_{comida}_{dia}"]
                    for dia in DIAS_SEMANA
                    if f"{prefijo}_{comida}_{dia}" in campos_por_identificador
                ]
                if campos_fila:
                    filas.append(
                        {
                            "titulo": comida.capitalize(),
                            "campos": campos_fila,
                        }
                    )
            secciones.append(
                {
                    "titulo": titulo,
                    "tipo": "matriz",
                    "identificador": identificador,
                    "filas": filas,
                }
            )
        return secciones

    @staticmethod
    def _formatear_valor_campo(informe, field):
        value = getattr(informe, field.name)

        if field.choices:
            display_value = getattr(informe, f"get_{field.name}_display", None)
            if display_value:
                try:
                    return display_value()
                except Exception:
                    pass

        if isinstance(value, bool):
            return "Sí" if value else "No"

        if isinstance(value, datetime):
            if timezone.is_aware(value):
                value = timezone.localtime(value)
            return value.strftime("%d/%m/%Y %H:%M")

        if isinstance(value, date):
            return value.strftime("%d/%m/%Y")

        return value

    @staticmethod
    def preparar_informe_para_creacion(instance, admision_id, action=None):
        """Inicializa un informe técnico nuevo según la acción (borrador/finalizado)."""
        try:
            instance.admision_id = admision_id
            if action == "draft":
                instance.estado_formulario = "borrador"
                instance.estado = "Iniciado"
            else:
                instance.estado_formulario = "finalizado"
                instance.estado = "Para revision"
        except Exception:
            logger.exception(
                "Error en preparar_informe_para_creacion",
                extra={"admision_pk": admision_id},
            )

    @staticmethod
    def get_informe_por_tipo_y_pk(tipo, pk):
        try:
            return get_object_or_404(InformeTecnico, tipo=tipo, pk=pk)
        except Exception:
            logger.exception(
                "Error en get_informe_por_tipo_y_pk",
                extra={"tipo": tipo, "informe_pk": pk},
            )
            return None

    @classmethod
    def _configurar_documento_template(cls, documento):
        """Configura la hoja base compartida por editor y DOCX dinámico."""

        seccion = documento.sections[0]
        seccion.page_width = Mm(cls.ANCHO_PAGINA_TEMPLATE_MM)
        seccion.page_height = Mm(cls.ALTO_PAGINA_TEMPLATE_MM)
        seccion.top_margin = Mm(cls.MARGEN_PAGINA_TEMPLATE_MM)
        seccion.right_margin = Mm(cls.MARGEN_PAGINA_TEMPLATE_MM)
        seccion.bottom_margin = Mm(cls.MARGEN_PAGINA_TEMPLATE_MM)
        seccion.left_margin = Mm(cls.MARGEN_PAGINA_TEMPLATE_MM)

        estilo_normal = documento.styles["Normal"]
        estilo_normal.font.name = cls.FUENTE_BASE_TEMPLATE
        estilo_normal.font.size = Pt(cls.TAMANIO_FUENTE_BASE_TEMPLATE_PT)

    @classmethod
    def _aplicar_estilo_tablas_templates(cls, documento):
        """Aplica una grilla legible a las tablas de templates dinámicos.

        htmldocx crea la estructura de la tabla pero no traslada los estilos CSS
        del editor al documento Word. El formato se aplica sobre OOXML para que
        el resultado no dependa de estilos HTML que el formulario no persiste.
        """

        for tabla in documento.tables:
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER  # pylint: disable=no-member
            cls._aplicar_bordes_tabla_template(tabla)
            for indice_fila, fila in enumerate(tabla.rows):
                es_encabezado = indice_fila == 0 and cls._fila_es_encabezado(fila)
                for celda in fila.cells:
                    celda.vertical_alignment = (  # pylint: disable=no-member
                        WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    )
                    cls._aplicar_margenes_celda_template(celda)
                    for parrafo in celda.paragraphs:
                        parrafo.paragraph_format.space_after = 0
                        parrafo.paragraph_format.space_before = 0
                    if es_encabezado:
                        cls._aplicar_fondo_encabezado_template(celda)

    @staticmethod
    def _fila_es_encabezado(fila):
        """Reconoce los <th> ya convertidos por htmldocx como texto en negrita."""

        return all(
            any(
                run.bold
                for parrafo in celda.paragraphs
                for run in parrafo.runs
                if run.text.strip()
            )
            for celda in fila.cells
        )

    @classmethod
    def _aplicar_bordes_tabla_template(cls, tabla):
        propiedades_tabla = tabla._tbl.tblPr  # pylint: disable=protected-access
        bordes = propiedades_tabla.first_child_found_in("w:tblBorders")
        if bordes is None:
            bordes = OxmlElement("w:tblBorders")
            propiedades_tabla.append(bordes)

        for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
            borde = bordes.find(qn(f"w:{lado}"))
            if borde is None:
                borde = OxmlElement(f"w:{lado}")
                bordes.append(borde)
            borde.set(qn("w:val"), "single")
            borde.set(qn("w:sz"), "6")
            borde.set(qn("w:space"), "0")
            borde.set(qn("w:color"), cls.COLOR_BORDE_TABLA_TEMPLATE)

    @staticmethod
    def _aplicar_margenes_celda_template(celda):
        propiedades_celda = (
            celda._tc.get_or_add_tcPr()
        )  # pylint: disable=protected-access
        margenes = propiedades_celda.first_child_found_in("w:tcMar")
        if margenes is None:
            margenes = OxmlElement("w:tcMar")
            propiedades_celda.append(margenes)

        for lado, valor in (
            ("top", "80"),
            ("start", "110"),
            ("bottom", "80"),
            ("end", "110"),
        ):
            margen = margenes.find(qn(f"w:{lado}"))
            if margen is None:
                margen = OxmlElement(f"w:{lado}")
                margenes.append(margen)
            margen.set(qn("w:w"), valor)
            margen.set(qn("w:type"), "dxa")

    @classmethod
    def _aplicar_fondo_encabezado_template(cls, celda):
        propiedades_celda = (
            celda._tc.get_or_add_tcPr()
        )  # pylint: disable=protected-access
        sombreado = propiedades_celda.find(qn("w:shd"))
        if sombreado is None:
            sombreado = OxmlElement("w:shd")
            propiedades_celda.append(sombreado)
        sombreado.set(qn("w:val"), "clear")
        sombreado.set(qn("w:color"), "auto")
        sombreado.set(qn("w:fill"), cls.COLOR_ENCABEZADO_TABLA_TEMPLATE)

    @staticmethod
    def actualizar_estado_informe(informe, nuevo_estado, tipo=None):
        try:
            informe.estado = nuevo_estado
            informe.save()

            # Actualizar estado de admisión según el nuevo estado del informe
            from ..admisiones_service import AdmisionService

            if nuevo_estado == "A subsanar":
                AdmisionService.actualizar_estado_admision(
                    informe.admision, "subsanar_informe"
                )
            elif nuevo_estado == "Validado":
                AdmisionService.actualizar_estado_admision(
                    informe.admision, "aprobar_informe_tecnico"
                )
                # No generar nuevos archivos - el DOCX del técnico es el final
        except Exception as e:
            logger.exception(
                f"Error en actualizar_estado_informe: {str(e)}",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            raise

    @staticmethod
    def _normalizar_tipo_admision(admision):
        tipo = getattr(admision, "tipo", None)
        normalizado = (tipo or "").strip().lower()
        return normalizado or "incorporacion"

    @staticmethod
    def generar_docx_con_template(informe, template_name=None):
        """Genera DOCX usando template con docxtpl"""
        try:
            # Seleccionar template basado en tipo de admisión e informe
            if not template_name:
                admision_tipo = InformeService._normalizar_tipo_admision(
                    informe.admision
                )
                informe_tipo = informe.tipo
                template_name = (
                    f"{admision_tipo}_docx_informe_tecnico_{informe_tipo}.docx"
                )

            context = AdmisionesContextService.preparar_contexto_informe_tecnico(
                informe
            )
            logger.debug(
                "Generando DOCX para informe %s con template %s (tipo %s)",
                informe.id,
                template_name,
                informe.tipo,
            )

            return DocumentTemplateService.generar_docx(template_name, context)
        except Exception as e:
            logger.exception(
                "No se pudo procesar template DOCX para informe %s: %s",
                getattr(informe, "pk", None),
                str(e),
            )
            return None

    @staticmethod
    def generar_docx_con_version_publicada(informe, version):
        """Renderiza una versión publicada del Gestor de templates a DOCX."""

        try:
            contexto = AdmisionesContextService.preparar_contexto_informe_tecnico(
                informe
            )
            contexto.update(
                {
                    "admision": informe.admision,
                    "comedor": getattr(informe.admision, "comedor", None),
                }
            )
            engine = Engine(debug=False, builtins=[], libraries={})
            contenido_renderizado = engine.from_string(version.contenido_html).render(
                Context(contexto, autoescape=True)
            )
            return InformeService._generate_docx_content(
                contenido_renderizado,
                getattr(informe, "pk", None),
                estilizar_tablas_templates=True,
            )
        except TemplateSyntaxError as error:
            logger.warning(
                "La versión de template %s tiene sintaxis inválida: %s",
                getattr(version, "pk", None),
                error,
            )
        except Exception:
            logger.exception(
                "No se pudo renderizar la versión de template %s",
                getattr(version, "pk", None),
            )
        return None

    @staticmethod
    def generar_docx_vista_previa(informe, version):
        """Genera un DOCX temporal con marca de agua sin persistirlo."""

        docx_content = InformeService.generar_docx_con_version_publicada(
            informe,
            version,
        )
        if not docx_content:
            return None
        try:
            docx_content.seek(0)
            documento = Document(BytesIO(docx_content.read()))
            marca_agua = (
                "<w:pict {namespaces}>"
                '<v:shape id="SISOCPreviewWatermark" '
                'o:spid="_x0000_s1025" type="#_x0000_t136" '
                'style="position:absolute;margin-left:0;margin-top:0;'
                "width:468pt;height:117pt;rotation:315;z-index:-251654144;"
                'mso-position-horizontal:center;mso-position-vertical:center" '
                'fillcolor="#d9d9d9" stroked="f">'
                '<v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt" '
                'string="VISTA PREVIA — DOCUMENTO NO VÁLIDO"/>'
                "</v:shape></w:pict>"
            ).format(namespaces=nsdecls("w", "v", "o"))
            for seccion in documento.sections:
                seccion.header.is_linked_to_previous = False
                parrafo = seccion.header.paragraphs[0]
                parrafo._p.append(parse_xml(marca_agua))

            buffer = BytesIO()
            documento.save(buffer)
            buffer.seek(0)
            return ContentFile(buffer.getvalue(), name="vista-previa.docx")
        except Exception:
            logger.exception(
                "No se pudo aplicar la marca de agua de vista previa",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return None

    @staticmethod
    def generar_y_guardar_pdf(informe, tipo):
        """
        Genera y guarda PDF y DOCX del informe técnico
        """
        from weasyprint import HTML

        try:
            context = {
                "informe": informe,
                "texto_comidas": generar_texto_comidas(informe),
            }

            # Mapear tipos de informe para templates
            admision_tipo = InformeService._normalizar_tipo_admision(informe.admision)
            informe_tipo_map = {
                "base": "base",
                "juridico": "juridico",
                "juridico eclesiastico": "juridico",
            }
            informe_tipo = informe_tipo_map.get(informe.tipo, "base")

            pdf_template = f"admisiones/pdf/{admision_tipo}_pdf_informe_tecnico_{informe_tipo}.html"
            docx_template = f"admisiones/docx/{admision_tipo}_docx_informe_tecnico_{informe_tipo}.html"

            logger.debug(
                "Generando archivos para informe %s (PDF: %s, DOCX: %s)",
                informe.id,
                pdf_template,
                docx_template,
            )

            # Generar PDF
            try:
                html_pdf = render_to_string(pdf_template, context)
                if not html_pdf.strip():
                    raise ValueError(f"Template PDF vacío: {pdf_template}")

                pdf_bytes = HTML(
                    string=html_pdf, base_url=InformeService._get_base_url()
                ).write_pdf()
                if not pdf_bytes:
                    raise ValueError("WeasyPrint no generó contenido PDF")

            except Exception as e:
                logger.error("Error generando PDF: %s", str(e))
                raise

            # Generar DOCX HTML template
            try:
                html_docx = render_to_string(docx_template, context)
                if not html_docx.strip():
                    html_docx = html_pdf
            except Exception as e:
                logger.warning("Error con template DOCX HTML: %s", str(e))
                html_docx = html_pdf

            # Generar DOCX con docxtpl
            docx_content = None
            try:
                logger.info(
                    f"Intentando generar DOCX con template para informe {informe.id}"
                )
                docx_buffer = InformeService.generar_docx_con_template(informe)
                if docx_buffer:
                    logger.info("DOCX generado exitosamente con template docxtpl")
                    docx_content = ContentFile(docx_buffer.getvalue(), name="tmp.docx")
                else:
                    raise ValueError("Template DOCX retornó None")
            except Exception as e:
                logger.warning("Template DOCX falló: %s, usando fallback HTML", str(e))
                docx_content = InformeService._generate_docx_content(
                    html_docx, getattr(informe, "pk", None)
                )

            # Guardar archivos - PDF final (sin prefijo "borrador")
            base_filename = (
                slugify(f"{tipo}-informe-{informe.id}") or f"informe-{informe.id}"
            )
            pdf_content = ContentFile(pdf_bytes, name=f"{base_filename}.pdf")

            defaults = {
                "tipo": tipo,
                "informe_id": informe.id,
                "comedor": getattr(informe.admision, "comedor", None),
                "archivo": pdf_content,
            }

            if docx_content:
                docx_content.name = f"{base_filename}.docx"
                defaults["archivo_docx"] = docx_content
                logger.debug(
                    "Archivos PDF y DOCX preparados para guardar para informe %s",
                    informe.id,
                )
            else:
                logger.debug(
                    "Solo se guardará PDF para informe %s, generación DOCX falló",
                    informe.id,
                )

            pdf_obj, created = InformeTecnicoPDF.objects.update_or_create(
                admision=informe.admision, defaults=defaults
            )

            action = "creado" if created else "actualizado"
            logger.info(
                "InformeTecnicoPDF %s exitosamente para informe %s",
                action,
                informe.id,
            )

            return pdf_obj

        except Exception as e:
            logger.exception(
                f"Error crítico en generar_y_guardar_pdf: {str(e)}",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            raise

    @staticmethod
    def get_informe_create_context(admision_id, tipo):
        try:
            admision = get_object_or_404(
                Admision.objects.select_related("comedor"), pk=admision_id
            )
            return {
                "tipo": tipo,
                "admision": admision,
                "comedor": admision.comedor,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_create_context",
                extra={"admision_pk": admision_id, "tipo": tipo},
            )
            return {}

    @staticmethod
    def get_informe_update_context(informe, tipo):
        try:
            campos_a_subsanar_db = CampoASubsanar.objects.filter(
                informe=informe
            ).values_list("campo", flat=True)

            # Incluir tanto nombres de campo como verbose names para el template
            campos_a_subsanar = list(campos_a_subsanar_db)
            field_to_verbose = {
                field.name: field.verbose_name for field in informe._meta.fields
            }

            # Agregar también los verbose names para mostrar en la interfaz
            for campo in campos_a_subsanar_db:
                verbose_name = field_to_verbose.get(campo, campo)
                if verbose_name not in campos_a_subsanar:
                    campos_a_subsanar.append(verbose_name)

            try:
                observacion = ObservacionGeneralInforme.objects.get(informe=informe)
            except ObservacionGeneralInforme.DoesNotExist:
                observacion = None

            return {
                "tipo": tipo,
                "admision": informe.admision,
                "comedor": informe.admision.comedor,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "campos_a_subsanar": campos_a_subsanar,
                "observacion": observacion,
            }
        except Exception:
            logger.exception(
                "Error en get_informe_update_context",
                extra={
                    "informe_pk": getattr(informe, "pk", None),
                    "tipo": tipo,
                },
            )
            return {}

    @staticmethod
    @transaction.atomic
    def guardar_informe(form, admision, es_creacion=False, action=None, usuario=None):
        """Guarda el informe técnico."""
        try:
            from ..admisiones_service import AdmisionService
            from ..templates_informe_tecnico_service import (
                PlantillaInformeTecnicoService,
            )

            if (
                action == "submit"
                and not AdmisionService._todos_obligatorios_tienen_archivos(admision)
            ):
                return {
                    "success": False,
                    "error": "No puede finalizar el informe técnico si faltan documentos obligatorios por cargar.",
                }
            if (
                action == "submit"
                and not AdmisionService._todos_obligatorios_aceptados(admision)
            ):
                return {
                    "success": False,
                    "error": "No puede finalizar el informe técnico si hay documentos obligatorios sin validar.",
                }

            if es_creacion:
                existente = (
                    InformeTecnico.objects.filter(
                        admision=admision,
                        tipo=form.instance.tipo,
                    )
                    .order_by("-id")
                    .first()
                )
                if existente:
                    if existente.estado == "Validado":
                        return {
                            "success": False,
                            "error": "El informe técnico ya fue validado y no puede editarse.",
                        }
                    form.instance.pk = existente.pk
                    form.instance._state.adding = False
                    es_creacion = False

            publicacion = None
            error_template = None
            if action == "submit":
                publicacion, error_template = (
                    PlantillaInformeTecnicoService.resolver_publicacion_para_admision(
                        admision
                    )
                )

            usar_fallback_heredado = (
                PlantillaInformeTecnicoService.es_configuracion_faltante(
                    error_template
                )
            )
            accion_para_estado = (
                "draft" if error_template and not usar_fallback_heredado else action
            )

            if es_creacion:
                InformeService.preparar_informe_para_creacion(
                    form.instance,
                    admision.id,
                    accion_para_estado,
                )
            else:
                InformeService.verificar_estado_para_revision(
                    form.instance,
                    accion_para_estado,
                )

            informe = form.save(commit=False)
            informe.admision = admision

            # Set creado_por and modificado_por fields
            if usuario:
                if es_creacion:
                    informe.creado_por = usuario
                informe.modificado_por = usuario

            # Limpiar observaciones de subsanación cuando se finaliza el informe
            if action == "submit" and informe.observaciones_subsanacion:
                informe.observaciones_subsanacion = None

            informe.save()
            if hasattr(form, "save_m2m"):
                form.save_m2m()

            if error_template and not usar_fallback_heredado:
                if es_creacion:
                    AdmisionService.actualizar_estado_admision(
                        admision,
                        "iniciar_informe_tecnico",
                    )
                return {
                    "success": False,
                    "error": (
                        f"{error_template} Se guardó el Informe Técnico como borrador; "
                        "puede corregir sus validaciones e intentar nuevamente."
                    ),
                }

            if action == "submit" and informe.estado_formulario == "finalizado":
                resultado_docx = InformeService.generar_docx_borrador(
                    informe,
                    publicacion,
                )
                if resultado_docx:
                    # Solo actualizar estado si el DOCX se generó exitosamente
                    AdmisionService.congelar_documentacion_organizacional(
                        admision, usuario
                    )
                    admision.estado_admision = "informe_tecnico_finalizado"
                    admision.save()
                else:
                    transaction.set_rollback(True)
                    return {
                        "success": False,
                        "error": "No se pudo generar el DOCX borrador. El Informe Técnico permanece editable.",
                    }

            # Actualizar estado de admisión según la acción
            if es_creacion and accion_para_estado != "submit":
                AdmisionService.actualizar_estado_admision(
                    admision, "iniciar_informe_tecnico"
                )

            return {"success": True, "informe": informe}
        except Exception as e:
            transaction.set_rollback(True)
            logger.exception(
                "Error en guardar_informe",
                extra={"admision_pk": getattr(admision, "pk", None)},
            )
            return {"success": False, "error": str(e)}

    @staticmethod
    def get_context_informe_detail(informe, tipo):
        try:
            pdf_filter = {
                "admision": informe.admision,
                "tipo": tipo,
                "informe_id": informe.id,
            }
            pdf_final = (
                InformeTecnicoPDF.objects.filter(**pdf_filter).first()
                if informe.estado == "Validado"
                else None
            )
            pdf_borrador = (
                InformeTecnicoPDF.objects.filter(**pdf_filter).first()
                if informe.estado_formulario == "finalizado"
                and informe.estado != "Validado"
                else None
            )

            return {
                "tipo": tipo,
                "admision": informe.admision,
                "campos": InformeService.get_campos_visibles_informe(informe),
                "pdf": pdf_final,
                "pdf_borrador": pdf_borrador,
            }
        except Exception:
            logger.exception(
                "Error en get_context_informe_detail",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )
            return {}

    @staticmethod
    def procesar_revision_informe(request, tipo, informe):
        try:
            nuevo_estado = request.POST.get("estado")
            if nuevo_estado not in ["A subsanar", "Validado"]:
                return

            InformeService.actualizar_estado_informe(informe, nuevo_estado, tipo)

            if nuevo_estado == "A subsanar":
                campos_a_subsanar = request.POST.getlist("campos_a_subsanar")
                observacion = request.POST.get("observacion", "").strip()

                CampoASubsanar.objects.filter(informe=informe).delete()

                verbose_to_field = {
                    field.verbose_name: field.name for field in informe._meta.fields
                }
                for campo in campos_a_subsanar:
                    field_name = verbose_to_field.get(campo, campo)
                    CampoASubsanar.objects.create(informe=informe, campo=field_name)

                # Guardar observaciones de subsanación en el informe
                informe.observaciones_subsanacion = observacion
                informe.save()

                obs_obj, _ = ObservacionGeneralInforme.objects.get_or_create(
                    informe=informe
                )
                obs_obj.texto = observacion
                obs_obj.save()
            else:
                CampoASubsanar.objects.filter(informe=informe).delete()
                ObservacionGeneralInforme.objects.filter(informe=informe).delete()
                # Limpiar observaciones de subsanación
                informe.observaciones_subsanacion = None
                informe.save()
        except Exception:
            logger.exception(
                "Error en procesar_revision_informe",
                extra={"informe_pk": getattr(informe, "pk", None), "tipo": tipo},
            )

    @staticmethod
    def guardar_campos_complementarios(informe_tecnico, campos_dict, usuario):
        """
        Guarda los campos modificados como un solo conjunto de cambios.
        Siempre actualiza el informe complementario existente o crea uno nuevo.
        """
        try:
            informe, created = InformeComplementario.objects.get_or_create(
                admision=informe_tecnico.admision,
                defaults={
                    "informe_tecnico": informe_tecnico,
                    "creado_por": usuario,
                    "estado": "borrador",
                },
            )

            InformeComplementarioCampos.objects.filter(
                informe_complementario=informe
            ).delete()

            for campo, valor in campos_dict.items():
                InformeComplementarioCampos.objects.create(
                    campo=campo, value=valor, informe_complementario=informe
                )

            return informe

        except Exception:
            logger.exception(
                "Error en guardar_campos_complementarios",
                extra={"informe_tecnico_pk": getattr(informe_tecnico, "pk", None)},
            )
            return None

    @staticmethod
    def generar_y_guardar_pdf_complementario(informe_complementario):
        try:
            with transaction.atomic():
                informe = informe_complementario.informe_tecnico
                campos_modificados = InformeComplementarioCampos.objects.filter(
                    informe_complementario=informe_complementario
                )

                verbose_to_field = {
                    field.verbose_name.lower().strip(): field.name
                    for field in informe._meta.fields
                }

                campos_actualizados_detalle = []

                for c in campos_modificados:
                    key = c.campo.lower().strip()
                    field_name = (
                        c.campo
                        if hasattr(informe, c.campo)
                        else verbose_to_field.get(key)
                    )
                    if not field_name:
                        continue

                    field = informe._meta.get_field(field_name)
                    valor_original = getattr(informe, field_name, None)
                    nuevo_valor = c.value

                    if field.get_internal_type() in [
                        "IntegerField",
                        "PositiveIntegerField",
                    ]:
                        try:
                            nuevo_valor = int(nuevo_valor) if nuevo_valor else 0
                        except ValueError:
                            nuevo_valor = None
                    elif field.get_internal_type() == "DateField":
                        from django.utils.dateparse import parse_date

                        parsed = parse_date(str(nuevo_valor))
                        if parsed:
                            nuevo_valor = parsed

                    setattr(informe, field_name, nuevo_valor)
                    campos_actualizados_detalle.append(
                        {
                            "campo": field.verbose_name or field.name,
                            "valor_anterior": valor_original,
                            "valor_nuevo": nuevo_valor,
                        }
                    )

                if campos_actualizados_detalle:
                    informe.save()

                usuario = informe_complementario.creado_por
                creado_por = usuario.username if usuario else "—"

                context = {
                    "informe": informe,
                    "texto_comidas": generar_texto_comidas(informe),
                    "campos_actualizados": campos_actualizados_detalle,
                    "fecha": getattr(informe_complementario, "creado", None),
                    "creado_por": creado_por,
                }

                from weasyprint import HTML

                pdf_template = "admisiones/pdf/informe_tecnico_complementario.html"
                html_pdf = render_to_string(pdf_template, context)
                pdf_bytes = HTML(
                    string=html_pdf, base_url=InformeService._get_base_url()
                ).write_pdf()

                docx_content = DocumentTemplateService.generar_docx(
                    template_name="informe_tecnico_complementario.docx",
                    context=context,
                    app_name="admisiones",
                )

                base_filename = slugify(
                    f"informe-complementario-{informe.id}-admision-{informe.admision_id}-{informe.tipo}"
                )

                defaults = {
                    "admision": informe.admision,
                    "tipo": informe.tipo,
                    "archivo": ContentFile(pdf_bytes, name=f"{base_filename}.pdf"),
                }

                if docx_content:
                    docx_content.seek(0)
                    defaults["archivo_docx"] = ContentFile(
                        docx_content.read(), name=f"{base_filename}.docx"
                    )

                InformeTecnicoComplementarioPDF.objects.update_or_create(
                    informe_complementario=informe_complementario,
                    defaults=defaults,
                )

                return True

        except Exception:
            import logging

            logging.exception("Error en generar_y_guardar_pdf_complementario")
            return None

    @staticmethod
    def generar_docx_borrador(informe, publicacion):
        """Genera el DOCX borrador desde la versión publicada aplicable."""
        try:
            if publicacion is None:
                logger.info(
                    "No hay template dinámico publicado para el informe %s; "
                    "se utiliza el generador heredado durante la transición.",
                    getattr(informe, "pk", None),
                )
                return InformeService._generar_docx_borrador_heredado(informe)
            if not getattr(publicacion, "version", None):
                logger.error(
                    "No se puede generar el DOCX del informe %s sin una publicación aplicable.",
                    getattr(informe, "pk", None),
                )
                return None

            docx_content = InformeService.generar_docx_con_version_publicada(
                informe,
                publicacion.version,
            )
            if not docx_content:
                return None

            # Guardar DOCX borrador
            base_filename = slugify(
                f"informe-{informe.id}-admision-{informe.admision_id}-{informe.tipo}-borrador"
            )
            docx_content.name = f"{base_filename}.docx"

            # Solo crear/actualizar si no existe un PDF final (validado)
            if informe.estado != "Validado":
                pdf_borrador, created = InformeTecnicoPDF.objects.update_or_create(
                    admision=informe.admision,
                    defaults={
                        "tipo": informe.tipo,
                        "informe_id": informe.id,
                        "comedor": informe.admision.comedor,
                        "archivo_docx": docx_content,
                        "plantilla_informe_tecnico": publicacion.plantilla,
                        "version_plantilla_informe_tecnico": publicacion.version,
                    },
                )

                # Actualizar estado del informe
                informe.estado = "Docx generado"
                informe.save()

                logger.info(
                    "DOCX borrador generado para informe %s, estado actualizado",
                    informe.id,
                )

                return pdf_borrador
            else:
                logger.info(
                    "No se genera DOCX borrador porque ya existe PDF final validado para informe %s",
                    informe.id,
                )
                return None

        except Exception:
            logger.exception(
                "Error en generar_docx_borrador",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return None

    @staticmethod
    def _generar_docx_borrador_heredado(informe):
        """Conserva la generación vigente hasta completar la configuración inicial."""

        try:
            docx_buffer = InformeService.generar_docx_con_template(informe)
            if docx_buffer:
                return ContentFile(docx_buffer.getvalue(), name="tmp.docx")
            raise ValueError("Template DOCX retornó None")
        except Exception as error:
            logger.warning(
                "Template DOCX heredado falló para informe %s: %s; usando fallback HTML",
                getattr(informe, "pk", None),
                error,
            )

        context = {
            "informe": informe,
            "texto_comidas": generar_texto_comidas(informe),
        }
        admision_tipo = InformeService._normalizar_tipo_admision(informe.admision)
        informe_tipo_map = {
            "base": "base",
            "juridico": "juridico",
            "juridico eclesiastico": "juridico",
        }
        informe_tipo = informe_tipo_map.get(informe.tipo, "base")
        docx_template = (
            f"admisiones/docx/{admision_tipo}_docx_informe_tecnico_{informe_tipo}.html"
        )
        try:
            html_docx = render_to_string(docx_template, context)
            return InformeService._generate_docx_content(
                html_docx,
                getattr(informe, "pk", None),
            )
        except Exception:
            logger.exception(
                "Fallback HTML heredado falló para informe %s",
                getattr(informe, "pk", None),
            )
            return None

    @staticmethod
    def subir_docx_editado(informe, archivo_docx, usuario=None):
        """Maneja la subida del DOCX editado por el técnico"""
        try:
            # Obtener o crear el registro PDF
            pdf_obj, created = InformeTecnicoPDF.objects.get_or_create(
                admision=informe.admision,
                defaults={
                    "tipo": informe.tipo,
                    "informe_id": informe.id,
                    "comedor": informe.admision.comedor,
                },
            )

            # Pisar el DOCX borrador con el editado
            base_filename = slugify(
                f"informe-{informe.id}-admision-{informe.admision_id}-{informe.tipo}-final"
            )
            archivo_docx.name = f"{base_filename}.docx"
            pdf_obj.archivo_docx_editado = archivo_docx
            if not created:
                pdf_obj.tipo = informe.tipo
                pdf_obj.informe_id = informe.id
                pdf_obj.comedor = informe.admision.comedor
            pdf_obj.save(
                update_fields=[
                    "archivo_docx_editado",
                    "tipo",
                    "informe_id",
                    "comedor",
                ]
            )

            # Actualizar estado del informe
            informe.estado = "Docx editado"
            informe.save()

            # Actualizar estado de admisión - ahora pasa a revisión
            from ..admisiones_service import AdmisionService

            AdmisionService.actualizar_estado_admision(
                informe.admision, "enviar_informe_revision"
            )

            logger.info(
                "DOCX editado subido exitosamente para informe %s",
                informe.id,
            )

            return pdf_obj

        except Exception:
            logger.exception(
                "Error en subir_docx_editado",
                extra={"informe_pk": getattr(informe, "pk", None)},
            )
            return None

    @staticmethod
    def obtener_cambios_complementarios_texto(informe_complementario):
        """Obtiene los cambios del complementario como texto para mostrar en la interfaz"""
        try:
            return list(
                InformeComplementarioCampos.objects.filter(
                    informe_complementario=informe_complementario
                )
            )
        except Exception:
            logger.exception(
                "Error en obtener_cambios_complementarios_texto",
                extra={
                    "informe_complementario_pk": getattr(
                        informe_complementario, "pk", None
                    )
                },
            )
            return []
